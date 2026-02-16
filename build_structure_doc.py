# -*- coding: utf-8 -*-
"""Generate a Word document of project structure (format similar to data0204.docx)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def tree_lines(root: Path, prefix: str = "", skip: set = None) -> list[str]:
    skip = skip or {".git", "__pycache__", ".venv", "venv"}
    lines = []
    try:
        entries = sorted(root.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower()))
    except OSError:
        return lines
    for i, p in enumerate(entries):
        if p.name in skip or p.name.startswith(".") and p.name != ".gitignore":
            continue
        is_last = i == len(entries) - 1
        conn = "    " if is_last else "|   "
        sym = "    " if is_last else "|   "
        name = p.name
        if p.is_dir():
            lines.append(prefix + "|   " + name + "/")
            lines.extend(tree_lines(p, prefix + conn, skip))
        else:
            lines.append(prefix + "|   " + name)
    return lines

def main():
    root = Path(__file__).resolve().parent
    doc = Document()
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project structure: sepsis-trajectory-python-data")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    # Directory structure
    doc.add_paragraph("Directory structure", style="Heading 1")
    tree = ["sepsis-trajectory-python-data/"]
    tree.append("|   README.md")
    tree.append("|   requirements.txt")
    tree.append("|   LICENSE")
    tree.append("|   .gitignore")
    tree.append("|   sepsis_trajectory.py")
    tree.append("|   src/")
    for f in ["__init__.py", "io_utils.py", "synthetic_data.py", "trajectory.py", "features.py",
              "modeling.py", "survival.py", "explain.py", "plots.py"]:
        tree.append("|   |   " + f)
    tree.append("|   data/")
    tree.append("|   |   sample_data.csv")
    tree.append("|   |   sample_data.xlsx")
    tree.append("|   |   synthetic_sample_data.csv")
    tree.append("|   models/")
    tree.append("|   |   pipeline.pkl")
    tree.append("|   |   feature_names.json")
    tree.append("|   |   split_indices.json")
    tree.append("|   results/")
    tree.append("|   |   bic_selection.png")
    tree.append("|   |   figure2_*.png, figure3_*.png, figure4_km.png")
    tree.append("|   |   cox_28d_summary.csv, model_metrics.json, run_summary.json")
    doc.add_paragraph("\n".join(tree), style="Normal")
    doc.add_paragraph()
    doc.add_paragraph("---")
    doc.add_paragraph()

    # 0) Run
    doc.add_paragraph("0) How to run", style="Heading 1")
    doc.add_paragraph("pip install -r requirements.txt", style="Normal")
    doc.add_paragraph("python sepsis_trajectory.py", style="Normal")
    doc.add_paragraph()
    doc.add_paragraph("After running, the following will be generated/overwritten:")
    doc.add_paragraph("- models/pipeline.pkl, models/feature_names.json, models/split_indices.json")
    doc.add_paragraph("- results/bic_selection.png, results/figure2_*.png, results/figure3_*.png, results/figure4_km.png")
    doc.add_paragraph("- results/cox_28d_summary.csv, results/model_metrics.json, results/run_summary.json")
    doc.add_paragraph("- results/feature_importance.csv (if SHAP/explain step is run)")
    doc.add_paragraph()
    doc.add_paragraph("---")
    doc.add_paragraph()

    # 1) requirements.txt
    doc.add_paragraph("1) requirements.txt", style="Heading 1")
    req_path = root / "requirements.txt"
    if req_path.exists():
        doc.add_paragraph(req_path.read_text(encoding="utf-8", errors="replace"), style="Normal")
    doc.add_paragraph()
    doc.add_paragraph("---")
    doc.add_paragraph()

    # 2) .gitignore
    doc.add_paragraph("2) .gitignore", style="Heading 1")
    gi_path = root / ".gitignore"
    if gi_path.exists():
        doc.add_paragraph(gi_path.read_text(encoding="utf-8", errors="replace"), style="Normal")
    doc.add_paragraph()
    doc.add_paragraph("---")
    doc.add_paragraph()

    # 3) README.md (abbreviated)
    doc.add_paragraph("3) README.md (overview)", style="Heading 1")
    readme_path = root / "README.md"
    if readme_path.exists():
        text = readme_path.read_text(encoding="utf-8", errors="replace")
        if len(text) > 2000:
            text = text[:2000] + "\n\n... (truncated)"
        doc.add_paragraph(text, style="Normal")
    doc.add_paragraph()
    doc.add_paragraph("---")
    doc.add_paragraph()

    # 4) LICENSE
    doc.add_paragraph("4) LICENSE", style="Heading 1")
    lic_path = root / "LICENSE"
    if lic_path.exists():
        text = lic_path.read_text(encoding="utf-8", errors="replace")
        if len(text) > 1500:
            text = text[:1500] + "\n\n... (truncated)"
        doc.add_paragraph(text, style="Normal")
    doc.add_paragraph()
    doc.add_paragraph("---")
    doc.add_paragraph()

    # 5) Main script summary
    doc.add_paragraph("5) sepsis_trajectory.py (main script summary)", style="Heading 1")
    main_path = root / "sepsis_trajectory.py"
    if main_path.exists():
        text = main_path.read_text(encoding="utf-8", errors="replace")
        doc.add_paragraph(text[:2500] + "\n\n... (truncated)" if len(text) > 2500 else text, style="Normal")
    doc.add_paragraph()
    doc.add_paragraph("---")
    doc.add_paragraph()

    # 6) src/ modules list
    doc.add_paragraph("6) src/ modules", style="Heading 1")
    doc.add_paragraph("__init__.py, io_utils.py, synthetic_data.py, trajectory.py, features.py,")
    doc.add_paragraph("modeling.py, survival.py, explain.py, plots.py")
    doc.add_paragraph()

    out_path = root.parent / "sepsis-trajectory-python-data-structure.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    main()
